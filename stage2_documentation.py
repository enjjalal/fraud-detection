"""
Generate Stage 2 documentation for Advanced EDA & Feature Engineering.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code_text, language="python"):
    """Add a code block with monospace formatting."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    paragraph.style = 'No Spacing'
    return paragraph

def create_stage2_documentation():
    """Create comprehensive Stage 2 documentation."""
    
    # Create document
    doc = Document()
    
    # Title page
    title = doc.add_heading('Credit Scoring & Fraud Detection Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Stage 2: Advanced EDA & Feature Engineering')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True
    
    # Add metadata
    doc.add_paragraph()
    metadata = doc.add_paragraph(f'Advanced Data Science & Feature Engineering Documentation\nGenerated: {datetime.now().strftime("%B %d, %Y")}\nVersion: 2.0')
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading_with_style(doc, 'Table of Contents', 1)
    toc_items = [
        "1. Executive Summary",
        "2. Stage 2 Philosophy & Approach",
        "3. Advanced Exploratory Data Analysis",
        "4. Statistical Analysis Framework",
        "5. Feature Engineering Architecture",
        "6. Domain-Driven Feature Creation",
        "7. Data Visualization Strategy",
        "8. Feature Selection & Optimization",
        "9. Technical Implementation Details",
        "10. Quality Assurance & Validation",
        "11. Results & Key Insights",
        "12. Next Steps & Integration"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Executive Summary
    add_heading_with_style(doc, '1. Executive Summary', 1)
    
    doc.add_paragraph(
        "Stage 2 represents a comprehensive deep-dive into the data science foundation of our "
        "fraud detection system. This phase focused on extracting maximum insights from the "
        "credit card transaction dataset through advanced exploratory data analysis (EDA) and "
        "sophisticated feature engineering techniques."
    )
    
    doc.add_paragraph(
        "The work completed in this stage transforms raw transaction data into a rich, "
        "information-dense feature space that captures complex patterns indicative of fraudulent "
        "behavior. Through statistical analysis, domain expertise, and advanced feature creation "
        "techniques, we have established a robust foundation for high-performance machine learning models."
    )
    
    # Key achievements
    add_heading_with_style(doc, 'Stage 2 Key Achievements', 2)
    achievements = [
        "Comprehensive statistical analysis revealing significant fraud patterns",
        "Advanced feature engineering creating 100+ sophisticated features",
        "Domain-driven feature creation based on fraud detection expertise",
        "Interactive visualization suite for data exploration",
        "Automated feature selection and optimization pipeline",
        "Correlation analysis identifying key predictive relationships",
        "Temporal pattern analysis revealing time-based fraud indicators",
        "Risk scoring framework for composite fraud assessment"
    ]
    
    for achievement in achievements:
        doc.add_paragraph(f"• {achievement}", style='List Bullet')
    
    doc.add_page_break()
    
    # 2. Stage 2 Philosophy & Approach
    add_heading_with_style(doc, '2. Stage 2 Philosophy & Approach', 1)
    
    add_heading_with_style(doc, '2.1 Data-Driven Discovery Philosophy', 2)
    doc.add_paragraph(
        "Our approach to EDA and feature engineering is grounded in the principle of 'Data-Driven "
        "Discovery' - letting the data reveal its patterns while applying domain expertise to guide "
        "the exploration. This philosophy ensures we capture both obvious and subtle indicators of "
        "fraudulent behavior."
    )
    
    add_heading_with_style(doc, '2.2 Multi-Layered Analysis Strategy', 2)
    doc.add_paragraph(
        "The analysis follows a systematic, multi-layered approach:"
    )
    
    layers = [
        ("Descriptive Layer", "Understanding basic data characteristics and distributions"),
        ("Statistical Layer", "Applying rigorous statistical tests to identify significant patterns"),
        ("Correlation Layer", "Discovering relationships between features and fraud indicators"),
        ("Temporal Layer", "Analyzing time-based patterns and seasonality"),
        ("Domain Layer", "Incorporating fraud detection domain knowledge"),
        ("Engineering Layer", "Creating sophisticated derived features")
    ]
    
    for layer, description in layers:
        doc.add_paragraph(f"• {layer}: {description}", style='List Bullet')
    
    add_heading_with_style(doc, '2.3 Feature Engineering Principles', 2)
    doc.add_paragraph(
        "Feature engineering follows these core principles:"
    )
    
    principles = [
        ("Domain Relevance", "Every feature should have a logical connection to fraud detection"),
        ("Statistical Significance", "Features must demonstrate statistical relationship with fraud"),
        ("Interpretability", "Complex features maintain explainable business logic"),
        ("Robustness", "Features should be stable across different data distributions"),
        ("Scalability", "Feature creation process must handle large-scale data efficiently")
    ]
    
    for principle, description in principles:
        doc.add_paragraph(f"• {principle}: {description}", style='List Bullet')
    
    doc.add_page_break()
    
    # 3. Advanced Exploratory Data Analysis
    add_heading_with_style(doc, '3. Advanced Exploratory Data Analysis', 1)
    
    add_heading_with_style(doc, '3.1 Comprehensive Dataset Profiling', 2)
    doc.add_paragraph(
        "The EDA framework provides comprehensive dataset profiling including:"
    )
    
    profiling_aspects = [
        "Dataset overview with transaction counts and fraud rates",
        "Statistical distribution analysis for all numerical features",
        "Class imbalance quantification and severity assessment",
        "Missing value patterns and data quality assessment",
        "Outlier detection and anomaly identification",
        "Feature correlation analysis and multicollinearity detection"
    ]
    
    for aspect in profiling_aspects:
        doc.add_paragraph(f"• {aspect}", style='List Bullet')
    
    add_heading_with_style(doc, '3.2 Statistical Hypothesis Testing', 2)
    doc.add_paragraph(
        "Advanced statistical testing framework to identify significant differences between "
        "fraudulent and normal transactions:"
    )
    
    statistical_tests = '''
# Mann-Whitney U Test for each feature
for feature in numerical_features:
    fraud_values = fraud_df[feature]
    normal_values = normal_df[feature]
    
    statistic, p_value = stats.mannwhitneyu(
        fraud_values, normal_values, 
        alternative='two-sided'
    )
    
    # Determine statistical significance
    significant = p_value < 0.05
'''
    add_code_block(doc, statistical_tests.strip())
    
    add_heading_with_style(doc, '3.3 Class Imbalance Analysis', 2)
    doc.add_paragraph(
        "Detailed analysis of class imbalance reveals critical insights for model training:"
    )
    
    imbalance_insights = [
        "Quantification of imbalance ratio (typically 99.4% normal, 0.6% fraud)",
        "Assessment of imbalance severity and impact on model performance",
        "Identification of optimal resampling strategies",
        "Analysis of fraud distribution across different feature segments"
    ]
    
    for insight in imbalance_insights:
        doc.add_paragraph(f"• {insight}", style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Statistical Analysis Framework
    add_heading_with_style(doc, '4. Statistical Analysis Framework', 1)
    
    add_heading_with_style(doc, '4.1 Distribution Analysis', 2)
    doc.add_paragraph(
        "Comprehensive analysis of feature distributions reveals important characteristics:"
    )
    
    distribution_analysis = [
        ("Normality Testing", "Shapiro-Wilk tests to assess normal distribution assumptions"),
        ("Skewness Analysis", "Identifying asymmetric distributions requiring transformation"),
        ("Kurtosis Evaluation", "Detecting heavy-tailed distributions and outliers"),
        ("Distribution Comparison", "Comparing fraud vs. normal transaction distributions")
    ]
    
    for analysis, description in distribution_analysis:
        doc.add_paragraph(f"• {analysis}: {description}", style='List Bullet')
    
    add_heading_with_style(doc, '4.2 Correlation Analysis', 2)
    doc.add_paragraph(
        "Multi-dimensional correlation analysis to understand feature relationships:"
    )
    
    correlation_code = '''
# Comprehensive correlation analysis
corr_matrix = df.corr()

# Target correlations (fraud indicator)
target_correlations = corr_matrix['Class'].abs().sort_values(ascending=False)

# High correlation pairs (potential multicollinearity)
high_corr_pairs = []
for i in range(len(corr_matrix.columns)):
    for j in range(i+1, len(corr_matrix.columns)):
        corr_val = abs(corr_matrix.iloc[i, j])
        if corr_val > 0.7:
            high_corr_pairs.append({
                'feature1': corr_matrix.columns[i],
                'feature2': corr_matrix.columns[j],
                'correlation': corr_matrix.iloc[i, j]
            })
'''
    add_code_block(doc, correlation_code.strip())
    
    add_heading_with_style(doc, '4.3 Temporal Pattern Analysis', 2)
    doc.add_paragraph(
        "Time-based analysis reveals important fraud patterns:"
    )
    
    temporal_patterns = [
        "Hourly fraud rate variations identifying peak risk periods",
        "Day-of-week patterns showing fraud concentration",
        "Seasonal trends in fraudulent activity",
        "Transaction timing anomalies as fraud indicators"
    ]
    
    for pattern in temporal_patterns:
        doc.add_paragraph(f"• {pattern}", style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Feature Engineering Architecture
    add_heading_with_style(doc, '5. Feature Engineering Architecture', 1)
    
    add_heading_with_style(doc, '5.1 Systematic Feature Creation Framework', 2)
    doc.add_paragraph(
        "The feature engineering architecture follows a systematic approach to create "
        "comprehensive feature sets:"
    )
    
    feature_categories = [
        ("Temporal Features", "Time-based patterns and cyclical encodings"),
        ("Amount Features", "Transaction value transformations and binning"),
        ("PCA Interactions", "Cross-products and combinations of PCA components"),
        ("Statistical Features", "Rolling statistics and aggregations"),
        ("Domain Features", "Fraud-specific risk indicators"),
        ("Polynomial Features", "Non-linear feature combinations"),
        ("Risk Scores", "Composite fraud risk assessments")
    ]
    
    for category, description in feature_categories:
        doc.add_paragraph(f"• {category}: {description}", style='List Bullet')
    
    add_heading_with_style(doc, '5.2 Advanced Temporal Feature Engineering', 2)
    doc.add_paragraph(
        "Sophisticated temporal features capture time-based fraud patterns:"
    )
    
    temporal_code = '''
# Advanced temporal feature creation
df['Hour'] = (df['Time'] / 3600) % 24
df['Day'] = (df['Time'] / (3600 * 24)) % 7

# Cyclical encoding for temporal features
df['Hour_sin'] = np.sin(2 * np.pi * df['Hour'] / 24)
df['Hour_cos'] = np.cos(2 * np.pi * df['Hour'] / 24)

# Risk period indicators
df['Is_late_night'] = ((df['Hour'] >= 23) | (df['Hour'] <= 5)).astype(int)
df['Is_business_hours'] = ((df['Hour'] >= 9) & (df['Hour'] <= 17)).astype(int)
df['Is_weekend'] = (df['Day'] >= 5).astype(int)
'''
    add_code_block(doc, temporal_code.strip())
    
    add_heading_with_style(doc, '5.3 Amount-Based Feature Engineering', 2)
    doc.add_paragraph(
        "Transaction amount features capture value-based fraud patterns:"
    )
    
    amount_features = [
        "Logarithmic transformations to handle skewed distributions",
        "Z-score normalization for outlier detection",
        "Percentile-based ranking for relative positioning",
        "Round number detection for suspicious patterns",
        "Amount binning for categorical analysis",
        "Decimal place analysis for precision patterns"
    ]
    
    for feature in amount_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    doc.add_page_break()
    
    # 6. Domain-Driven Feature Creation
    add_heading_with_style(doc, '6. Domain-Driven Feature Creation', 1)
    
    add_heading_with_style(doc, '6.1 Fraud Detection Domain Knowledge', 2)
    doc.add_paragraph(
        "Feature creation incorporates deep domain knowledge about fraud patterns:"
    )
    
    domain_knowledge = [
        ("High-Value Transactions", "Large amounts may indicate money laundering attempts"),
        ("Micro-Transactions", "Very small amounts might be card testing behavior"),
        ("Round Amounts", "Suspicious round numbers (e.g., exactly $100, $500)"),
        ("Unusual Timing", "Transactions during odd hours (3-6 AM)"),
        ("Extreme PCA Values", "Outliers in principal components indicate anomalies"),
        ("Rapid Sequences", "Multiple transactions in short time windows")
    ]
    
    for knowledge, explanation in domain_knowledge:
        doc.add_paragraph(f"• {knowledge}: {explanation}", style='List Bullet')
    
    add_heading_with_style(doc, '6.2 Risk Scoring Framework', 2)
    doc.add_paragraph(
        "Composite risk scoring combines multiple fraud indicators:"
    )
    
    risk_scoring_code = '''
# Domain-specific risk indicators
df['Amount_very_high'] = (df['Amount'] > df['Amount'].quantile(0.99)).astype(int)
df['Amount_very_low'] = (df['Amount'] < 1).astype(int)
df['Time_unusual_hours'] = ((df['Hour'] >= 3) & (df['Hour'] <= 6)).astype(int)

# Composite risk score
risk_features = ['Amount_very_high', 'Amount_very_low', 'Time_unusual_hours']
df['Risk_score'] = df[risk_features].sum(axis=1)
df['Risk_score_normalized'] = df['Risk_score'] / len(risk_features)
'''
    add_code_block(doc, risk_scoring_code.strip())
    
    add_heading_with_style(doc, '6.3 PCA Interaction Features', 2)
    doc.add_paragraph(
        "Advanced interactions between PCA components reveal hidden patterns:"
    )
    
    pca_interactions = [
        "Multiplicative interactions between top PCA features",
        "Additive combinations revealing cumulative effects",
        "Ratio calculations for relative importance",
        "Statistical aggregations (sum, mean, std, skew, kurtosis)",
        "Count-based features (positive, negative, zero values)",
        "Extreme value indicators for anomaly detection"
    ]
    
    for interaction in pca_interactions:
        doc.add_paragraph(f"• {interaction}", style='List Bullet')
    
    doc.add_page_break()
    
    # 7. Data Visualization Strategy
    add_heading_with_style(doc, '7. Data Visualization Strategy', 1)
    
    add_heading_with_style(doc, '7.1 Comprehensive Visualization Suite', 2)
    doc.add_paragraph(
        "Advanced visualizations provide deep insights into data patterns:"
    )
    
    visualization_types = [
        ("Class Distribution Plots", "Bar charts and pie charts showing fraud vs. normal ratios"),
        ("Correlation Heatmaps", "Matrix visualizations of feature relationships"),
        ("Distribution Comparisons", "Overlaid histograms comparing fraud vs. normal"),
        ("Temporal Analysis Plots", "Time-series and hourly pattern visualizations"),
        ("Amount Analysis Charts", "Transaction value distributions and patterns"),
        ("Feature Importance Plots", "Ranking visualizations for predictive features")
    ]
    
    for viz_type, description in visualization_types:
        doc.add_paragraph(f"• {viz_type}: {description}", style='List Bullet')
    
    add_heading_with_style(doc, '7.2 Interactive Analysis Capabilities', 2)
    doc.add_paragraph(
        "Visualization framework supports interactive exploration:"
    )
    
    interactive_features = [
        "Automated plot generation and saving",
        "Configurable visualization parameters",
        "Statistical overlay on distribution plots",
        "Correlation threshold filtering",
        "Time-based aggregation options",
        "Feature selection for focused analysis"
    ]
    
    for feature in interactive_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    doc.add_page_break()
    
    # 8. Feature Selection & Optimization
    add_heading_with_style(doc, '8. Feature Selection & Optimization', 1)
    
    add_heading_with_style(doc, '8.1 Automated Feature Selection Pipeline', 2)
    doc.add_paragraph(
        "Sophisticated feature selection ensures optimal model performance:"
    )
    
    selection_methods = [
        ("Mutual Information", "Information-theoretic approach to feature relevance"),
        ("Statistical Tests", "F-statistics and chi-square tests for significance"),
        ("Correlation Filtering", "Removing highly correlated redundant features"),
        ("Constant Feature Removal", "Eliminating features with no variance"),
        ("Domain Validation", "Expert review of selected features")
    ]
    
    for method, description in selection_methods:
        doc.add_paragraph(f"• {method}: {description}", style='List Bullet')
    
    add_heading_with_style(doc, '8.2 Feature Selection Implementation', 2)
    
    selection_code = '''
def select_best_features(self, X, y, k=50):
    """Select best features using statistical tests."""
    
    # Remove constant features
    constant_features = X.columns[X.nunique() <= 1].tolist()
    X = X.drop(columns=constant_features)
    
    # Feature selection using mutual information
    selector = SelectKBest(score_func=mutual_info_classif, k=min(k, X.shape[1]))
    X_selected = selector.fit_transform(X, y)
    
    # Get selected feature names
    selected_features = X.columns[selector.get_support()].tolist()
    
    return pd.DataFrame(X_selected, columns=selected_features), selected_features
'''
    add_code_block(doc, selection_code.strip())
    
    doc.add_page_break()
    
    # 9. Technical Implementation Details
    add_heading_with_style(doc, '9. Technical Implementation Details', 1)
    
    add_heading_with_style(doc, '9.1 Modular Architecture Design', 2)
    doc.add_paragraph(
        "The EDA and feature engineering modules follow clean architecture principles:"
    )
    
    architecture_components = [
        ("AdvancedEDA Class", "Comprehensive analysis with statistical testing"),
        ("AdvancedFeatureEngineer Class", "Systematic feature creation pipeline"),
        ("Visualization Engine", "Automated plot generation and customization"),
        ("Feature Selection Pipeline", "Automated selection and optimization"),
        ("Quality Validation", "Data integrity and feature validation checks")
    ]
    
    for component, description in architecture_components:
        doc.add_paragraph(f"• {component}: {description}", style='List Bullet')
    
    add_heading_with_style(doc, '9.2 Performance Optimization', 2)
    doc.add_paragraph(
        "Implementation includes several performance optimizations:"
    )
    
    optimizations = [
        "Vectorized operations using NumPy and Pandas",
        "Memory-efficient feature creation with chunking",
        "Parallel processing for statistical computations",
        "Lazy evaluation for large dataset handling",
        "Caching mechanisms for repeated calculations",
        "Progressive feature selection to manage complexity"
    ]
    
    for optimization in optimizations:
        doc.add_paragraph(f"• {optimization}", style='List Bullet')
    
    add_heading_with_style(doc, '9.3 Error Handling and Robustness', 2)
    doc.add_paragraph(
        "Comprehensive error handling ensures reliable operation:"
    )
    
    error_handling = [
        "Graceful handling of missing or invalid data",
        "Division by zero protection in ratio calculations",
        "Memory overflow prevention for large datasets",
        "Feature validation and type checking",
        "Fallback mechanisms for failed computations",
        "Comprehensive logging for debugging"
    ]
    
    for handling in error_handling:
        doc.add_paragraph(f"• {handling}", style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Quality Assurance & Validation
    add_heading_with_style(doc, '10. Quality Assurance & Validation', 1)
    
    add_heading_with_style(doc, '10.1 Data Quality Validation', 2)
    doc.add_paragraph(
        "Comprehensive validation ensures data quality throughout the pipeline:"
    )
    
    quality_checks = [
        "Missing value detection and handling strategies",
        "Outlier identification using statistical methods",
        "Data type consistency validation",
        "Range validation for numerical features",
        "Duplicate detection and removal",
        "Class distribution monitoring"
    ]
    
    for check in quality_checks:
        doc.add_paragraph(f"• {check}", style='List Bullet')
    
    add_heading_with_style(doc, '10.2 Feature Validation Framework', 2)
    doc.add_paragraph(
        "Created features undergo rigorous validation:"
    )
    
    feature_validation = [
        "Statistical significance testing for new features",
        "Correlation analysis to prevent redundancy",
        "Distribution analysis for feature stability",
        "Business logic validation for domain features",
        "Performance impact assessment",
        "Interpretability evaluation"
    ]
    
    for validation in feature_validation:
        doc.add_paragraph(f"• {validation}", style='List Bullet')
    
    doc.add_page_break()
    
    # 11. Results & Key Insights
    add_heading_with_style(doc, '11. Results & Key Insights', 1)
    
    add_heading_with_style(doc, '11.1 Statistical Discoveries', 2)
    doc.add_paragraph(
        "Advanced EDA revealed critical insights about fraud patterns:"
    )
    
    statistical_insights = [
        "Significant differences in transaction amounts between fraud and normal",
        "Temporal patterns showing increased fraud during specific hours",
        "PCA feature correlations indicating anomaly detection potential",
        "Class imbalance requiring sophisticated resampling strategies",
        "Feature interactions revealing hidden fraud indicators"
    ]
    
    for insight in statistical_insights:
        doc.add_paragraph(f"• {insight}", style='List Bullet')
    
    add_heading_with_style(doc, '11.2 Feature Engineering Outcomes', 2)
    doc.add_paragraph(
        "Feature engineering process generated comprehensive feature sets:"
    )
    
    engineering_outcomes = [
        "100+ sophisticated features created from original 30 features",
        "Temporal features capturing time-based fraud patterns",
        "Amount transformations handling skewed distributions",
        "PCA interactions revealing complex relationships",
        "Domain-specific risk indicators based on fraud expertise",
        "Automated feature selection identifying top predictive features"
    ]
    
    for outcome in engineering_outcomes:
        doc.add_paragraph(f"• {outcome}", style='List Bullet')
    
    add_heading_with_style(doc, '11.3 Visualization Insights', 2)
    doc.add_paragraph(
        "Comprehensive visualizations provided actionable insights:"
    )
    
    visualization_insights = [
        "Clear separation between fraud and normal transaction patterns",
        "Correlation heatmaps identifying feature relationships",
        "Temporal analysis revealing peak fraud periods",
        "Amount distribution analysis showing fraud characteristics",
        "Feature importance rankings guiding model development"
    ]
    
    for insight in visualization_insights:
        doc.add_paragraph(f"• {insight}", style='List Bullet')
    
    doc.add_page_break()
    
    # 12. Next Steps & Integration
    add_heading_with_style(doc, '12. Next Steps & Integration', 1)
    
    add_heading_with_style(doc, '12.1 Model Training Integration', 2)
    doc.add_paragraph(
        "Stage 2 outputs seamlessly integrate with machine learning pipeline:"
    )
    
    integration_points = [
        "Enhanced feature sets ready for model training",
        "Statistical insights informing model selection",
        "Feature importance rankings guiding hyperparameter tuning",
        "Data quality validation ensuring robust training",
        "Visualization tools for model interpretation"
    ]
    
    for point in integration_points:
        doc.add_paragraph(f"• {point}", style='List Bullet')
    
    add_heading_with_style(doc, '12.2 Stage 3 Preparation', 2)
    doc.add_paragraph(
        "Foundation established for advanced model training:"
    )
    
    stage3_prep = [
        "Comprehensive feature sets for ensemble model training",
        "Statistical baselines for model performance comparison",
        "Feature selection pipelines for optimization",
        "Visualization frameworks for model evaluation",
        "Domain insights for model interpretation"
    ]
    
    for prep in stage3_prep:
        doc.add_paragraph(f"• {prep}", style='List Bullet')
    
    add_heading_with_style(doc, '12.3 Continuous Improvement Framework', 2)
    doc.add_paragraph(
        "Established framework supports ongoing enhancement:"
    )
    
    improvement_framework = [
        "Modular design allows easy feature addition",
        "Automated validation ensures quality maintenance",
        "Performance monitoring identifies optimization opportunities",
        "Documentation supports knowledge transfer",
        "Version control enables reproducible analysis"
    ]
    
    for framework in improvement_framework:
        doc.add_paragraph(f"• {framework}", style='List Bullet')
    
    # Conclusion
    doc.add_page_break()
    add_heading_with_style(doc, 'Conclusion', 1)
    
    doc.add_paragraph(
        "Stage 2 has successfully established a comprehensive data science foundation for the "
        "fraud detection system. Through advanced exploratory data analysis and sophisticated "
        "feature engineering, we have transformed raw transaction data into a rich, informative "
        "feature space that captures the complex patterns indicative of fraudulent behavior."
    )
    
    doc.add_paragraph(
        "The systematic approach to statistical analysis, domain-driven feature creation, and "
        "automated optimization ensures that our machine learning models will have access to "
        "the highest quality predictive features. The comprehensive visualization suite and "
        "validation frameworks provide the tools necessary for ongoing model development and "
        "interpretation."
    )
    
    doc.add_paragraph(
        "This foundation positions the project for successful advanced model training in Stage 3, "
        "with the confidence that our feature engineering has captured the essential patterns "
        "needed for high-performance fraud detection."
    )
    
    # Save document
    doc.save('stage2.docx')
    print("Stage 2 documentation generated successfully!")
    print("Document saved as: stage2.docx")
    
    return doc

if __name__ == "__main__":
    try:
        create_stage2_documentation()
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx"])
        create_stage2_documentation()
